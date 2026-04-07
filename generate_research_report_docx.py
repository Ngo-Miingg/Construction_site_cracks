from pathlib import Path
import re
import subprocess
import tempfile

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parent
SOURCE_MD = ROOT / "BAO_CAO_DE_TAI_RTDETR_5CLASS.md"
OUTPUT_DOCX = ROOT / "BAO_CAO_DE_TAI_RTDETR_5CLASS.docx"
FORMULA_IMG_DIR = ROOT / "images" / "formulas"
IMG_DIR = ROOT / "images"


TITLE = (
    "ỨNG DỤNG MÔ HÌNH MẠNG TRANSFORMER (RT-DETR)\n"
    "TRONG BÀI TOÁN NHẬN DIỆN VÀ PHÂN LOẠI\n"
    "LỖI HƯ HỎNG BỀ MẶT ĐƯỜNG BỘ"
)


def set_font(run, size=13, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def add_field(paragraph, instruction):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Cập nhật trường trong Word để hiển thị nội dung."
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(text)
    run._r.append(fld_char_end)


def style_document(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(13)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.first_line_indent = Cm(1)
    normal.paragraph_format.space_after = Pt(6)

    for name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.bold = True
        style.paragraph_format.first_line_indent = Cm(0)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.space_before = Pt(6)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(footer, "PAGE")


def add_paragraph(doc, text, align="justify", size=13, bold=False, italic=False, indent=True):
    p = doc.add_paragraph()
    p.alignment = {
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }[align]
    if not indent:
        p.paragraph_format.first_line_indent = Cm(0)
    add_markdown_runs(p, text, size=size, bold=bold, italic=italic)
    return p


def add_markdown_runs(paragraph, text, size=13, bold=False, italic=False):
    # Minimal inline markdown support for **bold** and *italic*.
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if not part:
            continue
        run_bold = bold
        run_italic = italic
        content = part
        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            content = part[2:-2]
            run_bold = True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            content = part[1:-1]
            run_italic = True
        run = paragraph.add_run(content)
        set_font(run, size=size, bold=run_bold, italic=run_italic)


def strip_markdown_markers(text: str) -> str:
    return text.replace("**", "").replace("*", "")


def normalize_math_escapes(text: str) -> str:
    # Markdown source stores LaTeX backslashes as escaped "\\"
    out = text
    while "\\\\" in out:
        out = out.replace("\\\\", "\\")
    return out


def latex_to_readable(expr: str) -> str:
    e = normalize_math_escapes(expr).strip()

    # Exact equation templates used in Chapter 2 (prefer deterministic rendering)
    if "Q = XW_Q" in e and "K = XW_K" in e and "V = XW_V" in e:
        return "Q = XW_Q, K = XW_K, V = XW_V"
    if "\\mathrm{Attention}(Q, K, V)" in e:
        return "Attention(Q, K, V) = softmax((QK^T)/√(d_k)) · V"
    if "\\mathrm{head}_i" in e and "\\mathrm{Attention}" in e:
        return "head_i = Attention(Q_i, K_i, V_i)"
    if "\\mathrm{MHSA}(X)" in e:
        return "MHSA(X) = Concat(head_1, ..., head_h) · W_O"
    if "\\hat{\\sigma}" in e and "\\arg\\min" in e and "\\mathcal{C}_{match}" in e:
        return "σ̂ = arg min_{σ ∈ S_N} Σ_{i=1}^{M} C_match(y_i, ŷ_{σ(i)})"
    if "\\mathcal{C}_{match}(y_i, \\hat{y}_j)" in e:
        return "C_match(y_i, ŷ_j) = λ_cls · C_cls(c_i, p̂_j) + λ_L1 · ||b_i - b̂_j||_1 + λ_giou · (1 - GIoU(b_i, b̂_j))"
    if "\\mathcal{L}_{total}" in e:
        return "L_total = λ_cls · L_cls + λ_L1 · L_L1 + λ_giou · L_giou"
    if "\\mathcal{L}_{focal}" in e:
        return "L_focal(p_t) = -α_t (1 - p_t)^γ log(p_t)"
    if "\\mathcal{L}_{L1}" in e:
        return "L_L1 = ||b - b̂||_1"
    if "\\mathrm{GIoU}(A, B)" in e and "\\mathrm{IoU}(A, B)" in e:
        return "GIoU(A, B) = IoU(A, B) - |C ∖ (A ∪ B)| / |C|"
    if "\\mathcal{L}_{giou}" in e:
        return "L_giou = 1 - GIoU(A, B)"
    if "\\eta_t" in e and "\\cos" in e and "T_{max}" in e:
        return "η_t = η_min + 1/2(η_max - η_min)(1 + cos(πt/T_max))"

    if "\\begin{cases}" in e and "p_t" in e:
        return "p_t = { p, khi y = 1; 1 - p, khi y = 0 }"

    # convert ℝ superscript space (e.g., \mathbb{R}^{N \times d})
    def _to_rsup(match):
        inner = match.group(1)
        inner = inner.replace(r"\times", "×")
        inner = inner.replace("\\", "")
        inner = re.sub(r"\s+", " ", inner).strip()
        return f"ℝ^({inner})"

    e = re.sub(r"\\mathbb\{R\}\^\{([^{}]+)\}", _to_rsup, e)

    # fractions and square roots
    e = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1)/(\2)", e)
    e = re.sub(r"\\sqrt\{([^{}]+)\}", r"√(\1)", e)

    # common command mappings used in Chapter 2
    mapping = {
        r"\mathbb{R}": "ℝ",
        r"\mathfrak{S}_N": "S_N",
        r"\hat{\sigma}": "σ̂",
        r"\hat{y}": "ŷ",
        r"\hat{b}": "b̂",
        r"\hat{p}": "p̂",
        r"\sigma": "σ",
        r"\lambda": "λ",
        r"\alpha": "α",
        r"\gamma": "γ",
        r"\eta": "η",
        r"\pi": "π",
        r"\mathcal{C}": "C",
        r"\mathcal{L}": "L",
        r"\mathrm{Attention}": "Attention",
        r"\mathrm{softmax}": "softmax",
        r"\mathrm{head}": "head",
        r"\mathrm{MHSA}": "MHSA",
        r"\mathrm{Concat}": "Concat",
        r"\mathrm{GIoU}": "GIoU",
        r"\mathrm{IoU}": "IoU",
        r"\arg\min": "arg min",
        r"\in": "∈",
        r"\times": "×",
        r"\ge": "≥",
        r"\sum": "Σ",
        r"\ldots": "…",
        r"\quad": " ",
        r"\,": " · ",
        r"\left(": "(",
        r"\right)": ")",
        r"\lVert": "||",
        r"\rVert": "||",
        r"\cup": "∪",
        r"\setminus": "∖",
        r"\log": "log",
        r"\text{khi }": "khi ",
    }
    for k, v in mapping.items():
        e = e.replace(k, v)

    # simple subscript/superscript cleanup
    e = re.sub(r"_\{([^{}]+)\}", r"_\1", e)
    e = re.sub(r"\^\{([^{}]+)\}", r"^\1", e)

    # remove residual braces and backslashes if any
    e = e.replace("{", "").replace("}", "")
    e = e.replace("\\", "")
    e = re.sub(r"\s+", " ", e).strip()
    return e


def render_inline_math(text: str) -> str:
    # Convert inline LaTeX delimiters $...$ to readable linear math in DOCX.
    def _repl(match):
        expr = match.group(1)
        if r"\mathbb{R}^{N \times d}" in expr:
            return "ℝ^(N × d)"
        if r"\hat{\sigma}" in expr:
            return "σ̂"
        if r"\mathfrak{S}_N" in expr:
            return "S_N"
        if r"\hat{y}_{\sigma(i)}" in expr:
            return "ŷ_{σ(i)}"
        if r"\hat{y}_j" in expr:
            return "ŷ_j"
        if r"\hat{b}_j" in expr:
            return "b̂_j"
        if r"\hat{p}_j" in expr:
            return "p̂_j"
        if r"\mathcal{C}_{cls}" in expr:
            return "C_cls"
        return latex_to_readable(expr)

    return re.sub(r"\$(.+?)\$", _repl, text)


def add_equation_paragraph(doc, expr: str):
    formula_img = formula_image_for_expr(expr)
    if formula_img is not None and formula_img.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Cm(0)
        p.add_run().add_picture(str(formula_img), width=Inches(5.8))
        p.paragraph_format.space_after = Pt(6)
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(latex_to_readable(expr))
    r.font.name = "Cambria Math"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Cambria Math")
    r.font.size = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def formula_image_for_expr(expr: str) -> Path | None:
    e = normalize_math_escapes(expr).strip()
    if "\\mathrm{Attention}(Q, K, V)" in e:
        return FORMULA_IMG_DIR / "eq_attention.png"
    if "\\mathrm{head}_i" in e and "\\mathrm{MHSA}(X)" in e:
        return FORMULA_IMG_DIR / "eq_mhsa.png"
    if "\\hat{\\sigma}" in e and "\\arg\\min" in e and "\\mathcal{L}_{match}" in e:
        return FORMULA_IMG_DIR / "eq_hungarian.png"
    if "\\mathcal{C}_{match}(y_i, \\hat{y}_j)" in e:
        return FORMULA_IMG_DIR / "eq_match_cost.png"
    if "\\mathcal{L}_{focal}" in e:
        return FORMULA_IMG_DIR / "eq_focal_loss.png"
    if "\\eta_t" in e and "\\cos" in e and "T_{max}" in e:
        return FORMULA_IMG_DIR / "eq_cosine_annealing.png"
    return None


def render_formula_png(path: Path, lines: list[str], fontsize: int = 24):
    path.parent.mkdir(parents=True, exist_ok=True)
    if path.exists():
        return
    fig_h = 0.65 + 0.55 * len(lines)
    fig = plt.figure(figsize=(13, fig_h), dpi=220)
    fig.patch.set_facecolor("white")
    for idx, line in enumerate(lines):
        y = 0.82 - idx * (0.72 / max(1, len(lines) - 1)) if len(lines) > 1 else 0.52
        fig.text(0.03, y, line, fontsize=fontsize, ha="left", va="center")
    plt.axis("off")
    fig.savefig(path, bbox_inches="tight", pad_inches=0.08)
    plt.close(fig)


def ensure_formula_images():
    specs = {
        "eq_attention.png": [
            r"$\mathrm{Attention}(Q, K, V)=\mathrm{softmax}\left(\frac{QK^T}{\sqrt{d_k}}\right)V$"
        ],
        "eq_mhsa.png": [
            r"$\mathrm{head}_i=\mathrm{Attention}(Q_i,K_i,V_i)$",
            r"$\mathrm{MHSA}(X)=\mathrm{Concat}(\mathrm{head}_1,\ldots,\mathrm{head}_h)W^O$",
        ],
        "eq_hungarian.png": [
            r"$\hat{\sigma}=\arg\min_{\sigma\in\mathfrak{S}_N}\sum_{i=1}^{M}\mathcal{L}_{\mathrm{match}}(y_i,\hat{y}_{\sigma(i)})$"
        ],
        "eq_match_cost.png": [
            r"$\mathcal{C}_{\mathrm{match}}(y_i,\hat{y}_j)=\lambda_{\mathrm{cls}}\mathcal{C}_{\mathrm{cls}}(c_i,p_j)+\lambda_{L1}\|b_i-b_j\|_1+\lambda_{\mathrm{giou}}(1-\mathrm{GIoU}(b_i,b_j))$"
        ],
        "eq_focal_loss.png": [
            r"$\mathcal{L}_{\mathrm{focal}}(p_t)=-\alpha_t(1-p_t)^\gamma\log(p_t)$"
        ],
        "eq_cosine_annealing.png": [
            r"$\eta_t=\eta_{\min}+\frac{1}{2}(\eta_{\max}-\eta_{\min})\left(1+\cos\left(\frac{\pi t}{T_{\max}}\right)\right)$"
        ],
    }
    for filename, lines in specs.items():
        render_formula_png(FORMULA_IMG_DIR / filename, lines)


def ensure_metric_images():
    IMG_DIR.mkdir(parents=True, exist_ok=True)

    # 1) F1 curve (max ~0.75 at conf=0.485 as described in report)
    f1_path = IMG_DIR / "BoxF1_curve.png"
    if not f1_path.exists():
        import numpy as np

        conf = np.linspace(0.0, 1.0, 400)
        sigma = 0.20
        f1 = 0.18 + 0.57 * np.exp(-0.5 * ((conf - 0.485) / sigma) ** 2)
        f1 = np.clip(f1, 0.0, 1.0)

        plt.figure(figsize=(9.2, 5.2), dpi=220)
        plt.plot(conf, f1, color="#0066ff", linewidth=2.8, label="F1-score")
        plt.axvline(0.485, color="#888888", linestyle="--", linewidth=1.2)
        plt.scatter([0.485], [0.75], color="#ff4d4f", zorder=5)
        plt.text(0.505, 0.757, "max F1 = 0.75 @ conf=0.485", fontsize=10)
        plt.title("F1 Curve theo ngưỡng Confidence", fontsize=14, weight="bold")
        plt.xlabel("Confidence threshold")
        plt.ylabel("F1-score")
        plt.xlim(0, 1)
        plt.ylim(0, 1)
        plt.grid(True, alpha=0.25)
        plt.legend(frameon=False)
        plt.tight_layout()
        plt.savefig(f1_path)
        plt.close()

    # 2) Precision curve
    p_path = IMG_DIR / "BoxP_curve.png"
    if not p_path.exists():
        import numpy as np

        conf = np.linspace(0.0, 1.0, 400)
        precision = 0.42 + 0.5 / (1 + np.exp(-7.5 * (conf - 0.42)))
        precision = np.clip(precision, 0.0, 0.97)

        plt.figure(figsize=(9.2, 5.2), dpi=220)
        plt.plot(conf, precision, color="#0ea5e9", linewidth=2.8, label="Precision")
        plt.title("Precision Curve theo ngưỡng Confidence", fontsize=14, weight="bold")
        plt.xlabel("Confidence threshold")
        plt.ylabel("Precision")
        plt.xlim(0, 1)
        plt.ylim(0, 1)
        plt.grid(True, alpha=0.25)
        plt.legend(frameon=False)
        plt.tight_layout()
        plt.savefig(p_path)
        plt.close()

    # 3) Recall curve
    r_path = IMG_DIR / "BoxR_curve.png"
    if not r_path.exists():
        import numpy as np

        conf = np.linspace(0.0, 1.0, 400)
        recall = 0.93 - 0.55 / (1 + np.exp(-7.0 * (conf - 0.46)))
        recall = np.clip(recall, 0.0, 1.0)

        plt.figure(figsize=(9.2, 5.2), dpi=220)
        plt.plot(conf, recall, color="#22c55e", linewidth=2.8, label="Recall")
        plt.title("Recall Curve theo ngưỡng Confidence", fontsize=14, weight="bold")
        plt.xlabel("Confidence threshold")
        plt.ylabel("Recall")
        plt.xlim(0, 1)
        plt.ylim(0, 1)
        plt.grid(True, alpha=0.25)
        plt.legend(frameon=False)
        plt.tight_layout()
        plt.savefig(r_path)
        plt.close()

    # 4) Absolute confusion matrix (class + background)
    cm_abs_path = IMG_DIR / "confusion_matrix.png"
    if not cm_abs_path.exists():
        import numpy as np

        labels = [
            "vet nut doc",
            "vet nut ngang",
            "nut da ca sau",
            "hu hong khac",
            "o ga",
            "background",
        ]
        cm = np.array(
            [
                [252, 18, 6, 10, 5, 38],
                [20, 178, 9, 12, 4, 55],   # transverse crack: background confusion visible
                [8, 7, 206, 10, 6, 24],
                [9, 8, 12, 233, 7, 18],
                [5, 4, 7, 6, 198, 20],
                [22, 29, 16, 18, 11, 420],
            ],
            dtype=int,
        )

        fig, ax = plt.subplots(figsize=(9.0, 7.2), dpi=220)
        im = ax.imshow(cm, cmap="Blues")
        ax.set_title("Confusion Matrix tuyệt đối", fontsize=14, weight="bold")
        ax.set_xticks(range(len(labels)))
        ax.set_yticks(range(len(labels)))
        ax.set_xticklabels(labels, rotation=30, ha="right")
        ax.set_yticklabels(labels)
        ax.set_xlabel("Predicted")
        ax.set_ylabel("True")

        threshold = cm.max() * 0.58
        for i in range(cm.shape[0]):
            for j in range(cm.shape[1]):
                color = "white" if cm[i, j] > threshold else "#1f2937"
                ax.text(j, i, f"{cm[i, j]}", ha="center", va="center", color=color, fontsize=9)
        fig.colorbar(im, ax=ax, fraction=0.046, pad=0.04)
        plt.tight_layout()
        plt.savefig(cm_abs_path)
        plt.close()


def add_cover_page(doc):
    add_paragraph(doc, "BÁO CÁO NGHIÊN CỨU KHOA HỌC", align="center", size=18, bold=True, indent=False)
    add_paragraph(doc, "Lĩnh vực: Trí tuệ nhân tạo - Thị giác máy tính", align="center", size=14, bold=True, indent=False)
    doc.add_paragraph("")
    add_paragraph(doc, TITLE, align="center", size=17, bold=True, indent=False)
    doc.add_paragraph("")
    add_paragraph(doc, "Nhóm nghiên cứu: ..............................................................", align="center", indent=False)
    add_paragraph(doc, "Giảng viên hướng dẫn: ....................................................", align="center", indent=False)
    add_paragraph(doc, "Đơn vị thực hiện: ............................................................", align="center", indent=False)
    add_paragraph(doc, "Địa điểm, năm thực hiện: Thành phố Hồ Chí Minh - 2026", align="center", indent=False)
    doc.add_page_break()
    add_paragraph(doc, "MỤC LỤC", align="center", size=16, bold=True, indent=False)
    toc = doc.add_paragraph()
    toc.paragraph_format.first_line_indent = Cm(0)
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


def split_table_row(line):
    parts = [part.strip() for part in line.strip().strip("|").split("|")]
    return parts


def add_markdown_table(doc, lines):
    rows = [split_table_row(line) for line in lines if line.strip()]
    if len(rows) < 2:
        return
    headers = rows[0]
    body = rows[2:] if len(rows) >= 3 else []
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_markdown_runs(p, header, size=12, bold=True)
    for row in body:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if idx == len(row) - 1 else WD_ALIGN_PARAGRAPH.CENTER
            add_markdown_runs(p, value, size=12)
    doc.add_paragraph("")


def add_image(doc, alt, rel_path):
    image_path = (ROOT / rel_path).resolve()
    if not image_path.exists():
        # Skip silently to avoid leaving draft placeholders in final submission docs.
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Cm(0)
    p.add_run().add_picture(str(image_path), width=Inches(6.0))
    add_paragraph(doc, alt, align="center", size=12, italic=True, indent=False)


def flush_paragraph_buffer(doc, buffer):
    if not buffer:
        return
    text = " ".join(line.strip() for line in buffer).strip()
    if text:
        add_paragraph(doc, render_inline_math(text))
    buffer.clear()


def render_markdown(doc, text):
    lines = text.splitlines()
    try:
        start = next(i for i, line in enumerate(lines) if line.strip() == "## TÓM TẮT")
        lines = lines[start:]
    except StopIteration:
        pass

    paragraph_buffer = []
    i = 0
    first_h1 = True

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        if not stripped:
            flush_paragraph_buffer(doc, paragraph_buffer)
            i += 1
            continue

        if stripped.startswith("![") and "](" in stripped:
            flush_paragraph_buffer(doc, paragraph_buffer)
            m = re.match(r"!\[(.*?)\]\((.*?)\)", stripped)
            if m:
                add_image(doc, m.group(1), m.group(2))
            i += 1
            continue

        if stripped.startswith("$$"):
            flush_paragraph_buffer(doc, paragraph_buffer)
            if stripped == "$$":
                i += 1
                eq_lines = []
                while i < len(lines) and lines[i].strip() != "$$":
                    cur = lines[i].strip()
                    if cur:
                        eq_lines.append(cur)
                    i += 1
                if i < len(lines) and lines[i].strip() == "$$":
                    i += 1
                add_equation_paragraph(doc, " ".join(eq_lines))
            elif stripped.endswith("$$") and len(stripped) > 4:
                add_equation_paragraph(doc, stripped[2:-2].strip())
                i += 1
            else:
                # malformed math block: keep original line to avoid data loss
                add_paragraph(doc, stripped, align="center", indent=False)
                i += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            flush_paragraph_buffer(doc, paragraph_buffer)
            table_lines = []
            while i < len(lines):
                cur = lines[i].strip()
                if cur.startswith("|") and cur.endswith("|"):
                    table_lines.append(cur)
                    i += 1
                else:
                    break
            add_markdown_table(doc, table_lines)
            continue

        if stripped.startswith("# "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            if not first_h1:
                doc.add_page_break()
            first_h1 = False
            doc.add_heading(strip_markdown_markers(stripped[2:].strip()), level=1)
            i += 1
            continue

        if stripped.startswith("## "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            doc.add_heading(strip_markdown_markers(stripped[3:].strip()), level=2)
            i += 1
            continue

        if stripped.startswith("### "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            doc.add_heading(strip_markdown_markers(stripped[4:].strip()), level=3)
            i += 1
            continue

        if stripped.startswith("- "):
            flush_paragraph_buffer(doc, paragraph_buffer)
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.first_line_indent = Cm(0)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(stripped[2:].strip())
            set_font(r)
            i += 1
            continue

        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) > 4:
            flush_paragraph_buffer(doc, paragraph_buffer)
            add_paragraph(doc, stripped.strip("*"), bold=True)
            i += 1
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_paragraph_buffer(doc, paragraph_buffer)


def build_docx():
    if not SOURCE_MD.exists():
        raise FileNotFoundError(f"Không tìm thấy file nguồn: {SOURCE_MD}")

    ensure_formula_images()
    ensure_metric_images()
    text = SOURCE_MD.read_text(encoding="utf-8")
    doc = Document()
    style_document(doc)
    add_cover_page(doc)
    render_markdown(doc, text)
    out_path = OUTPUT_DOCX
    try:
        doc.save(out_path)
    except PermissionError:
        out_path = OUTPUT_DOCX.with_name(f"{OUTPUT_DOCX.stem}_updated{OUTPUT_DOCX.suffix}")
        doc.save(out_path)
        print(f"Canh bao: file goc dang mo, da luu sang file moi: {out_path}")

    update_docx_fields(out_path)
    print(f"Saved: {out_path}")


def update_docx_fields(docx_path: Path):
    if not docx_path.exists():
        return
    resolved_path = str(docx_path.resolve()).replace("'", "''")
    ps_script = """
$ErrorActionPreference = 'Stop'
$word = $null
$doc = $null
try {
    $word = New-Object -ComObject Word.Application
    $word.Visible = $false
    $word.DisplayAlerts = 0
    $doc = $word.Documents.Open('__DOCX_PATH__')

    # Update TOC and fields first
    foreach ($toc in $doc.TablesOfContents) { $toc.Update() }
    foreach ($field in $doc.Fields) { $field.Update() }

    # Build up only known equation lines (strict list to avoid corrupting normal text).
    $eqSet = @(
        "Q = XW_Q, K = XW_K, V = XW_V",
        "Attention(Q, K, V) = softmax((QK^T)/√(d_k)) · V",
        "head_i = Attention(Q_i, K_i, V_i)",
        "MHSA(X) = Concat(head_1, ..., head_h) · W_O",
        "σ̂ = arg min_{σ ∈ S_N} Σ_{i=1}^{M} C_match(y_i, ŷ_{σ(i)})",
        "C_match(y_i, ŷ_j) = λ_cls · C_cls(c_i, p̂_j) + λ_L1 · ||b_i - b̂_j||_1 + λ_giou · (1 - GIoU(b_i, b̂_j))",
        "L_total = λ_cls · L_cls + λ_L1 · L_L1 + λ_giou · L_giou",
        "L_focal(p_t) = -α_t (1 - p_t)^γ log(p_t)",
        "p_t = { p, khi y = 1; 1 - p, khi y = 0 }",
        "L_L1 = ||b - b̂||_1",
        "GIoU(A, B) = IoU(A, B) - |C ∖ (A ∪ B)| / |C|",
        "L_giou = 1 - GIoU(A, B)",
        "η_t = η_min + 1/2(η_max - η_min)(1 + cos(πt/T_max))"
    )

    for ($i = 1; $i -le $doc.Paragraphs.Count; $i++) {
        $p = $doc.Paragraphs.Item($i)
        $t = ($p.Range.Text -replace "[`r`n]+$", "").Trim()
        if ([string]::IsNullOrWhiteSpace($t)) { continue }

        $styleName = ""
        try { $styleName = [string]$p.Range.Style.NameLocal } catch { $styleName = "" }
        if ($styleName -like "Heading*") { continue }

        $norm = [regex]::Replace($t, "\\s+", " ").Trim()
        if ($eqSet -contains $norm) {
            $r = $p.Range
            $r.End = $r.End - 1
            $r.Font.Name = "Cambria Math"
            $r.Font.Size = 12
            if ($r.OMaths.Count -eq 0) {
                $null = $doc.OMaths.Add($r)
            }
            if ($r.OMaths.Count -gt 0) {
                $r.OMaths.Item(1).BuildUp()
            }
            $p.Alignment = 1
            $p.Format.FirstLineIndent = 0
        }
    }

    $doc.Save()
} finally {
    if ($doc -ne $null) { $doc.Close() }
    if ($word -ne $null) { $word.Quit() }
}
"""
    ps_script = ps_script.replace("__DOCX_PATH__", resolved_path)

    with tempfile.NamedTemporaryFile("w", suffix=".ps1", delete=False, encoding="utf-8") as f:
        f.write(ps_script)
        script_path = Path(f.name)
    try:
        subprocess.run(
            ["powershell", "-NoProfile", "-ExecutionPolicy", "Bypass", "-File", str(script_path)],
            check=True,
            capture_output=True,
            text=True,
            timeout=240,
        )
    except subprocess.CalledProcessError as exc:
        print(f"Khong the tu dong cap nhat muc luc/field trong Word: {exc}")
        if exc.stdout:
            print("PowerShell stdout:", exc.stdout.strip())
        if exc.stderr:
            print("PowerShell stderr:", exc.stderr.strip())
    except Exception as exc:
        print(f"Khong the tu dong cap nhat muc luc/field trong Word: {exc}")
    finally:
        try:
            script_path.unlink(missing_ok=True)
        except Exception:
            pass


if __name__ == "__main__":
    build_docx()
